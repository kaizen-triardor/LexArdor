"""Document template engine — 3-step pipeline: Parse → Label → Generalize.

Based on ChatGPT's approach to legal document templating:
1. PARSE: Split document into structural sections
2. LABEL: Identify what each section IS (zaglavlje, stranke, činjenični opis, zahtev...)
3. GENERALIZE: Replace concrete data with placeholders, keeping legal structure

This approach understands STRUCTURE first, then abstracts the variable parts.
"""
import re
import json
from core.transliterate import to_latin, to_cyrillic, detect_script

FIELD_TYPES = {
    "text": {"label": "Tekst", "validation": None},
    "jmbg": {"label": "JMBG", "validation": r"^\d{13}$"},
    "pib": {"label": "PIB", "validation": r"^\d{9}$"},
    "date": {"label": "Datum", "validation": r"^\d{1,2}\.\d{1,2}\.\d{4}\.?$"},
    "money": {"label": "Iznos", "validation": r"^[\d.,]+\s*(RSD|EUR|USD|din\.?|dinara)?$"},
    "address": {"label": "Adresa", "validation": None},
    "phone": {"label": "Telefon", "validation": r"^[\d\s\+\-()]+$"},
    "email": {"label": "Email", "validation": r"^[\w.+-]+@[\w-]+\.[\w.-]+$"},
    "multiline": {"label": "Duži tekst", "validation": None},
}

DOC_CATEGORIES = {
    "tuzba": "Tužba", "ugovor": "Ugovor", "resenje": "Rešenje",
    "punomocje": "Punomoćje", "dopis": "Dopis", "zalba": "Žalba",
    "podnesak": "Podnesak", "ostalo": "Ostalo",
}

# ═══════════════════════════════════════════════════════════════════════════════
# Step 0: Pre-processing
# ═══════════════════════════════════════════════════════════════════════════════

def _normalize_field_labels(fields: list[dict], target_script: str = "latin") -> list[dict]:
    for field in fields:
        label = field.get("label", "")
        if detect_script(label) in ("cyrillic", "mixed"):
            field["label"] = to_latin(label)
        name = field.get("name", "")
        name = to_latin(name) if detect_script(name) != "latin" else name
        name = re.sub(r"[^\w]", "_", name.lower()).strip("_")
        name = re.sub(r"_+", "_", name)
        field["name"] = name
    return fields


def _extract_first_document(text: str) -> str:
    """Extract only the first legal document from multi-document PDFs."""
    # Find first court/authority header
    start = 0
    for m in re.finditer(r'(?:OPŠTINSK|OSNOVN|VIŠEM|APELACION|PRIVREDN|VRHOVN|USTAVNOM|REPUBLIČKOM)\w*\s+SUDU', text, re.IGNORECASE):
        start = max(0, m.start() - 20)
        break

    # Find where the next document starts (look for another court header after 500+ chars)
    end = len(text)
    search_from = start + 500
    for m in re.finditer(r'(?:OPŠTINSK|OSNOVN|VIŠEM|APELACION|PRIVREDN)\w*\s+SUDU', text[search_from:], re.IGNORECASE):
        end = search_from + m.start()
        break

    result = text[start:min(end, start + 15000)]
    return result.strip()


# ═══════════════════════════════════════════════════════════════════════════════
# The 3-Step Pipeline: Parse → Label → Generalize
# ═══════════════════════════════════════════════════════════════════════════════

def analyze_document(text: str, llm_client) -> dict:
    """3-step pipeline to create a template from a filled legal document.

    Step 1: AI parses the document into labeled sections
    Step 2: AI identifies variable fields within each section
    Step 3: Build template by replacing variables with placeholders
    """
    first_doc = _extract_first_document(text)
    doc_text = first_doc[:6000] if len(first_doc) > 6000 else first_doc

    # ── STEP 1+2 Combined: Parse structure AND extract fields ──
    # One well-designed prompt that follows the ChatGPT approach:
    # "Structure > Content" — understand sections first, then find variables

    prompt = f"""Ti si pravni asistent koji pravi šablone (template) iz popunjenih pravnih dokumenata.

ZADATAK U 3 KORAKA:

KORAK 1 — STRUKTURA: Identifikuj sekcije dokumenta:
- zaglavlje (sud, mesto)
- stranke (tužilac, tuženi, podnosilac...)
- naslov (tip dokumenta, "radi čega")
- činjenični opis (šta se desilo)
- dokazi (ako postoje)
- pravni osnov (članovi zakona, ako postoje)
- zahtev (šta se traži od suda)
- završni deo (datum, potpis, mesto)

KORAK 2 — VARIJABLE: U svakoj sekciji pronađi SVE podatke koji bi se menjali za drugog klijenta:
- Imena lica (tužilac, tuženi, svedoci)
- Adrese
- JMBG, PIB brojeve
- Datume
- Novčane iznose
- Brojeve predmeta/ugovora
- Nazive sudova
- Opise činjenica i zahteva

KORAK 3 — TEMPLATE: Za svaku varijablu napravi placeholder.

PRAVILA:
- "name" polja piši u snake_case NA LATINICI (npr. ime_tuzioca)
- "label" piši na srpskoj LATINICI (npr. "Ime tuzioca")
- "example" mora biti TAČNA vrednost iz dokumenta
- Ako dokument sadrži više primera, koristi SAMO PRVI

ODGOVORI U OVOM JSON FORMATU (bez drugog teksta):
{{
  "document_type": "tuzba",
  "sections": ["zaglavlje", "stranke", "naslov", "cinjenicni_opis", "zahtev", "zavrsni_deo"],
  "fields": [
    {{"name": "naziv_suda", "label": "Naziv suda", "type": "text", "required": true, "section": "zaglavlje", "example": "Osnovnom sudu u Krusevcu"}},
    {{"name": "ime_tuzioca", "label": "Ime i prezime tuzioca", "type": "text", "required": true, "section": "stranke", "example": "Marinkovic Vera"}},
    {{"name": "adresa_tuzioca", "label": "Adresa tuzioca", "type": "address", "required": true, "section": "stranke", "example": "Partizanska 12, Krusevac"}},
    {{"name": "predmet_tuzbe", "label": "Predmet tuzbe (radi cega)", "type": "text", "required": true, "section": "naslov", "example": "ispraznjenje poslovnog prostora i naknada stete"}},
    {{"name": "cinjenicni_opis", "label": "Opis cinjenica", "type": "multiline", "required": true, "section": "cinjenicni_opis", "example": "Tuzene obavljaju frizersku delatnost..."}},
    {{"name": "iznos", "label": "Tuzeni iznos", "type": "money", "required": false, "section": "zahtev", "example": "14.850,00 dinara"}},
    {{"name": "datum_podnosenja", "label": "Datum podnosenja", "type": "date", "required": true, "section": "zavrsni_deo", "example": "15.09.2004"}}
  ]
}}

DOKUMENT ZA ANALIZU:
{doc_text}"""

    response = llm_client.generate(prompt, temperature=0.1, max_tokens=3000)
    result = _parse_ai_response(response)
    fields = result.get("fields", [])

    # Normalize labels to Latin
    fields = _normalize_field_labels(fields, "latin")

    # ── STEP 3: Build template body ──
    body_template = first_doc
    example_values = {}

    # Sort by example length descending (replace longer strings first)
    fields_sorted = sorted(fields, key=lambda f: len(f.get("example", "")), reverse=True)
    for field in fields_sorted:
        example = field.get("example", "")
        if not example:
            continue
        placeholder = "{{" + field["name"] + "}}"
        new_body = _fuzzy_replace(body_template, example, placeholder)
        if new_body != body_template:
            body_template = new_body
            example_values[field["name"]] = example

    return {
        "body_template": body_template,
        "fields": fields,
        "example_values": example_values,
        "document_type": result.get("document_type", "ostalo"),
        "sections": result.get("sections", []),
    }


# ═══════════════════════════════════════════════════════════════════════════════
# AI Response Parsing (robust, multi-strategy)
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_ai_response(response: str) -> dict:
    """Parse AI JSON response with multiple fallback strategies."""
    # Strategy 1: find outermost JSON object
    try:
        match = re.search(r'\{[\s\S]*\}', response)
        if match:
            raw = match.group()
            raw = raw.replace("'", '"')
            raw = re.sub(r',\s*}', '}', raw)
            raw = re.sub(r',\s*]', ']', raw)
            data = json.loads(raw)
            if data.get("fields"):
                return data
    except json.JSONDecodeError:
        pass

    # Strategy 2: find JSON array
    try:
        match = re.search(r'\[[\s\S]*\]', response)
        if match:
            raw = match.group().replace("'", '"')
            raw = re.sub(r',\s*]', ']', raw)
            fields = json.loads(raw)
            if fields:
                return {"fields": fields}
    except json.JSONDecodeError:
        pass

    # Strategy 3: regex extraction
    fields = []
    pattern = r'"name"\s*:\s*"([^"]+)".*?"label"\s*:\s*"([^"]+)".*?"type"\s*:\s*"([^"]+)".*?"example"\s*:\s*"([^"]*)"'
    for m in re.finditer(pattern, response, re.DOTALL):
        fields.append({
            "name": m.group(1), "label": m.group(2),
            "type": m.group(3), "required": True, "example": m.group(4),
        })
    return {"fields": fields}


def _fuzzy_replace(text: str, target: str, replacement: str) -> str:
    """Replace target in text, tolerating whitespace differences from PDF extraction."""
    if not target:
        return text
    if target in text:
        return text.replace(target, replacement, 1)
    # Normalize whitespace for matching
    norm_target = re.sub(r'\s+', r'\\s+', re.escape(target))
    match = re.search(norm_target, text)
    if match:
        return text[:match.start()] + replacement + text[match.end():]
    # Try core words only
    words = [w for w in target.split() if len(w) > 2]
    if len(words) >= 2:
        pattern = r'\s+'.join(re.escape(w) for w in words)
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return text[:match.start()] + replacement + text[match.end():]
    return text


# ═══════════════════════════════════════════════════════════════════════════════
# Fill, Validate, Smart-Fill
# ═══════════════════════════════════════════════════════════════════════════════

def fill_template(body_template: str, field_values: dict) -> str:
    """Replace {{placeholders}} with actual values."""
    result = body_template
    for name, value in field_values.items():
        result = result.replace("{{" + name + "}}", str(value))
    return result


def smart_fill_from_text(description: str, fields: list[dict], llm_client) -> dict:
    """AI extracts field values from a free-text case description."""
    field_list = "\n".join(
        f"- {f['name']} ({f['label']}, tip: {f['type']})"
        for f in fields
    )
    prompt = f"""Iz sledećeg opisa predmeta izvuci vrednosti za ova polja:

{field_list}

Odgovori ISKLJUČIVO u JSON formatu: {{"field_name": "vrednost", ...}}
Ako ne možeš da nađeš vrednost, stavi null.

OPIS PREDMETA:
{description}"""

    response = llm_client.generate(prompt, temperature=0.1, max_tokens=1000)
    try:
        match = re.search(r'\{[\s\S]*\}', response)
        if match:
            return json.loads(match.group())
    except json.JSONDecodeError:
        pass
    return {}


def validate_document(field_values: dict, fields: list[dict]) -> list[dict]:
    """Validate field values against field types."""
    issues = []
    for field in fields:
        name = field["name"]
        value = field_values.get(name, "")
        ftype = field.get("type", "text")
        required = field.get("required", False)
        label = field.get("label", name)

        if required and not value:
            issues.append({
                "field": name, "label": label, "level": "error",
                "message": f"{label} je obavezno polje",
            })
            continue
        if not value:
            continue
        type_info = FIELD_TYPES.get(ftype, {})
        pattern = type_info.get("validation")
        if pattern and not re.match(pattern, str(value)):
            issues.append({
                "field": name, "label": label, "level": "warning",
                "message": f"{label}: neispravan format za tip '{ftype}'",
            })
    return issues


# ═══════════════════════════════════════════════════════════════════════════════
# Export: DOCX and PDF/HTML
# ═══════════════════════════════════════════════════════════════════════════════

def generate_docx(filled_text: str, title: str = "") -> bytes:
    """Generate a properly formatted legal DOCX document."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    for line in filled_text.split('\n'):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Court/authority names → centered, bold
        if re.match(r'^(?:OPŠTINSK|OSNOVN|VIŠEM|APELACION|PRIVREDN|VRHOVN)', stripped, re.IGNORECASE):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(stripped).bold = True
        # Document title → centered, bold, larger
        elif re.match(r'^T\s*U\s*Ž\s*B\s*A|^TUŽBA|^UGOVOR|^REŠENJE|^PUNOMOĆJE|^ŽALBA|^P\s*R\s*E\s*S\s*U\s*D', stripped, re.IGNORECASE):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(14)
        # "radi ..." subtitle → centered, italic
        elif re.match(r'^radi\s', stripped, re.IGNORECASE):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(stripped).italic = True
        # Party lines → bold
        elif re.match(r'^(?:TUŽIL|TUŽEN|PODNOSIL|ZASTUPNIK|STRANKA|PUNOMOĆNIK)', stripped, re.IGNORECASE):
            p = doc.add_paragraph()
            p.add_run(stripped).bold = True
        # Section headers → centered, bold
        elif re.match(r'^(?:[IVX]+\.\s|PREDLOG|OBRAZLOŽENJE|TUŽBENI ZAHTEV|DOKAZI)', stripped, re.IGNORECASE):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(stripped).bold = True
            p.paragraph_format.space_before = Pt(12)
        # Date/signature → right-aligned
        elif re.match(r'^U\s+\w+,?\s+dana|^Tužilac|^Podnosilac|^_+', stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(stripped)
        # Regular paragraph → indented
        else:
            p = doc.add_paragraph()
            p.add_run(stripped)
            p.paragraph_format.first_line_indent = Cm(1)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_pdf_html(filled_text: str, title: str = "") -> str:
    """Generate formatted HTML for PDF export / preview."""
    text = filled_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    lines = text.split('\n')
    body_parts = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            body_parts.append('<div class="spacer"></div>')
            continue

        if re.match(r'^(?:OPŠTINSK|OSNOVN|VIŠEM|APELACION|PRIVREDN)', stripped, re.IGNORECASE):
            body_parts.append(f'<p class="center bold">{stripped}</p>')
        elif re.match(r'^T\s*U\s*Ž\s*B\s*A|^TUŽBA|^UGOVOR|^REŠENJE|^PUNOMOĆJE', stripped, re.IGNORECASE):
            body_parts.append(f'<h2 class="center">{stripped}</h2>')
        elif re.match(r'^radi\s', stripped, re.IGNORECASE):
            body_parts.append(f'<p class="center italic">{stripped}</p>')
        elif re.match(r'^(?:TUŽIL|TUŽEN|PODNOSIL|ZASTUPNIK)', stripped, re.IGNORECASE):
            body_parts.append(f'<p class="bold">{stripped}</p>')
        elif re.match(r'^(?:[IVX]+\.\s|PREDLOG|OBRAZLOŽENJE|TUŽBENI|DOKAZI)', stripped, re.IGNORECASE):
            body_parts.append(f'<h3 class="center">{stripped}</h3>')
        elif re.match(r'^U\s+\w+,?\s+dana|^Tužilac|^Podnosilac|^_+', stripped):
            body_parts.append(f'<p class="right">{stripped}</p>')
        else:
            body_parts.append(f'<p class="indent">{stripped}</p>')

    return f"""<!DOCTYPE html>
<html lang="sr"><head><meta charset="utf-8"><title>{title}</title>
<style>
@page {{ size: A4; margin: 2.5cm 2cm 2.5cm 3cm; }}
body {{ font-family: 'Times New Roman', serif; font-size: 12pt; line-height: 1.6; max-width: 700px; margin: 0 auto; padding: 40px; }}
h2 {{ font-size: 14pt; margin: 20px 0 10px; }} h3 {{ font-size: 12pt; margin: 16px 0 8px; }}
p {{ margin: 3px 0; }} .center {{ text-align: center; }} .right {{ text-align: right; }}
.bold {{ font-weight: bold; }} .italic {{ font-style: italic; }} .indent {{ text-indent: 1cm; }}
.spacer {{ height: 8px; }}
</style></head><body>{''.join(body_parts)}</body></html>"""
